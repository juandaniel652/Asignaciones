from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Crear un nuevo documento
doc = Document()

# ---- MODIFICAR ENCABEZADO ----
heading = doc.add_heading("Encabezado con Estilo", level=1)
run = heading.runs[0]
run.font.size = Pt(24)  # Tamaño de fuente grande
run.font.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)  # Color verde

# ---- MODIFICAR PÁRRAFO Y AÑADIR BORDE ----
p = doc.add_paragraph("Este es un párrafo con borde y fuente personalizada.")
run = p.runs[0]
run.font.size = Pt(14)  # Tamaño de fuente
run.font.italic = True
run.font.color.rgb = RGBColor(0, 0, 255)  # Color azul

# Agregar borde al párrafo (requiere manipular XML)
p_element = p._element
p_props = p_element.get_or_add_pPr()
borders = parse_xml(r'<w:pBdr {}><w:top w:val="single" w:sz="8" w:space="4" w:color="000000"/>'
                    r'<w:left w:val="single" w:sz="8" w:space="4" w:color="000000"/>'
                    r'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="000000"/>'
                    r'<w:right w:val="single" w:sz="8" w:space="4" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
p_props.append(borders)

# ---- AGREGAR UNA TABLA CON BORDES ----
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'  # Aplica bordes automáticamente

# Llenar la tabla con contenido
for row in table.rows:
    for cell in row.cells:
        cell.text = "Celda"
        for para in cell.paragraphs:
            para.alignment = 1  # Centrar texto en cada celda
            run = para.runs[0]
            run.font.size = Pt(12)  # Tamaño de fuente en la tabla

# Guardar el documento
doc.save("documento_bordes.docx")