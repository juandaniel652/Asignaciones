from docx import Document
from docx.shared import Pt, RGBColor

# Cargar el documento existente o crear uno nuevo
doc = Document()

# Agregar un encabezado con estilo modificado
heading = doc.add_heading("Encabezado Personalizado", level=1)
run = heading.runs[0]
run.font.size = Pt(20)  # Tamaño de fuente
run.font.bold = True  # Negrita
run.font.color.rgb = RGBColor(255, 0, 0)  # Color rojo

# Agregar un párrafo con estilo modificado
p = doc.add_paragraph("Este es un párrafo con estilo modificado.")
run = p.runs[0]
run.font.size = Pt(12)  # Tamaño de fuente
run.font.italic = True  # Cursiva
run.font.color.rgb = RGBColor(0, 0, 255)  # Color azul

# Guardar el documento
doc.save("documento_editado.docx")
