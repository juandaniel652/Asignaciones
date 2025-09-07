from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def widgets(subtitulos_y_parrafos):
    
    # Genera un documento de Word con varios subtítulos subrayados y sus respectivos párrafos
    
    doc = Document()


    for elementos in range (5) : 
        
        for subtitulo, parrafos in subtitulos_y_parrafos:

            # Crear subtítulo como encabezado
            subtitle = doc.add_paragraph()
            run = subtitle.add_run(subtitulo)
            run.font.size = Pt(18)  # Tamaño de fuente
            run.font.bold = True  # Negrita

            # Agregar subrayado inferior al subtítulo
            p_element = subtitle._element
            p_props = p_element.get_or_add_pPr()

            borders = parse_xml(r'''
            <w:pBdr {}>
                <w:bottom w:val="single" w:sz="24" w:space="4" w:color="000000"/>
            </w:pBdr>
            '''.format(nsdecls('w')))
            p_props.append(borders)

            # Agregar párrafos correspondientes a cada subtítulo
            for p_text in parrafos:
                p = doc.add_paragraph(p_text)
                run = p.runs[0]
                run.font.size = Pt(12)

            # Agregar un espacio extra entre secciones
            doc.add_paragraph()

    # Guardar documento
    doc.save("documento_unico.docx")
    print("Documento guardado: documento_unico.docx")

# Datos de prueba

for elementos in range (5) : 

    subtitulos_y_parrafos = [
        ("Primer Encabezado", ["Este es el primer párrafo.", "Este es el segundo párrafo."]),
        ("Segundo Encabezado", ["Aquí hay otro párrafo.", "Y aquí otro más."]),
        ("Tercer Encabezado", ["Más contenido en otro apartado.", "Último párrafo de ejemplo."])
    ]


# Generar documento con varios encabezados y párrafos
widgets(subtitulos_y_parrafos)
