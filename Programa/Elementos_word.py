from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def generar_documento_word(titulos_de_acomodadores_y_vigilancia, lista_de_los_parrafos, semanas_totales, cantidad_de_datos):

    documento = Document()
    for cantidad in range(cantidad_de_datos) : 

        titulo_de_semanas = documento.add_heading(f"{semanas_totales[cantidad]}", level = 1)
        estilo_de_titulo_de_semanas = titulo_de_semanas.runs[0]
        estilo_de_titulo_de_semanas.font.size = Pt(20)
        estilo_de_titulo_de_semanas.font.bold = True
        estilo_de_titulo_de_semanas.font.color.rgb = RGBColor(201, 201, 201) 

        edicion_del_titulo_a_partir_de_xml = titulo_de_semanas._element
        modificacion_de_las_propiedades_del_titulo = edicion_del_titulo_a_partir_de_xml.get_or_add_pPr()

        diseño_de_los_bordes_titulo = parse_xml(r'''
        <w:pBdr {}>
            <w:top w:val="single" w:sz="24" w:space="4" w:color="000000"/>
            <w:left w:val="single" w:sz="24" w:space="4" w:color="000000"/>
            <w:bottom w:val="24pt" w:space="4" w:color="000000"/>
            <w:right w:val="single" w:sz="24" w:space="4" w:color="000000"/>
        </w:pBdr>
        '''.format(nsdecls('w')))

        modificacion_de_las_propiedades_del_titulo.append(diseño_de_los_bordes_titulo)

        simulacion_de_degradado = parse_xml(r'<w:shd {} w:fill="303030"/>'.format(nsdecls('w')))
        modificacion_de_las_propiedades_del_titulo.append(simulacion_de_degradado)

        for largo_de_la_lista_parrafos, subtitulo in enumerate(titulos_de_acomodadores_y_vigilancia[cantidad]):

            parrafo_para_el_subtitulo = documento.add_paragraph()
            estilo_parrafo_para_el_subtitulo = parrafo_para_el_subtitulo.add_run(subtitulo)
            estilo_parrafo_para_el_subtitulo.font.size = Pt(18) 
            estilo_parrafo_para_el_subtitulo.font.bold = True  

            edicion_del_subtitulo_a_partir_de_xml = parrafo_para_el_subtitulo._element
            modificacion_de_las_propiedades_del_subtitulo = edicion_del_subtitulo_a_partir_de_xml.get_or_add_pPr()

            diseño_de_los_bordes_subtitulo = parse_xml(r'''
            <w:pBdr {}>
                <w:bottom w:val="single" w:sz="24" w:space="4" w:color="000000"/>
            </w:pBdr>
            '''.format(nsdecls('w')))

            modificacion_de_las_propiedades_del_subtitulo.append(diseño_de_los_bordes_subtitulo)

            if largo_de_la_lista_parrafos < len(lista_de_los_parrafos) :

                for textos_de_los_parrafos in range(len(lista_de_los_parrafos[cantidad][largo_de_la_lista_parrafos])): 

                    parrafos = documento.add_paragraph(lista_de_los_parrafos[cantidad][largo_de_la_lista_parrafos][textos_de_los_parrafos])
                    estilo_parrafo_para_el_subtitulo = parrafos.runs[0]
                    estilo_parrafo_para_el_subtitulo.font.size = Pt(13)
                    estilo_parrafo_para_el_subtitulo.font.italic = True

    return documento;    